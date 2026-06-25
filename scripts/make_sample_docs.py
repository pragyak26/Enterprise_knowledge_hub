"""Generate realistic sample company documents for testing the Knowledge Hub.

Creates .docx files (a supported upload type) under sample_docs/, including two
versions of the leave policy so you can try /ask and /compare.

Run:  python scripts/make_sample_docs.py
"""
import sys
from pathlib import Path

from docx import Document

OUT = Path(__file__).resolve().parent.parent / "sample_docs"


def write(name: str, title: str, paragraphs: list[str]) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for p in paragraphs:
        if p.startswith("# "):
            doc.add_heading(p[2:], level=1)
        else:
            doc.add_paragraph(p)
    path = OUT / name
    doc.save(path)
    print("wrote", path)


def main() -> int:
    OUT.mkdir(exist_ok=True)

    write(
        "LeavePolicy_v1.docx",
        "Leave Policy (v1)",
        [
            "# Annual Leave",
            "All full-time employees are entitled to 20 days of paid annual leave per calendar year. Unused leave may be carried over up to a maximum of 5 days.",
            "# Sick Leave",
            "Employees are entitled to 10 days of paid sick leave per year. A medical certificate is required for absences longer than 3 consecutive days.",
            "# Maternity Leave",
            "Employees are entitled to 120 days of paid maternity leave. Leave must commence no earlier than 30 days before the expected due date.",
            "# Paternity Leave",
            "Employees are entitled to 5 days of paid paternity leave, to be taken within 60 days of the child's birth.",
            "# Travel Reimbursement",
            "Business travel expenses are reimbursed up to a cap of 500 USD per trip. Original receipts must be submitted within 14 days.",
        ],
    )

    write(
        "LeavePolicy_v2.docx",
        "Leave Policy (v2)",
        [
            "# Annual Leave",
            "All full-time employees are entitled to 22 days of paid annual leave per calendar year. Unused leave may be carried over up to a maximum of 10 days.",
            "# Sick Leave",
            "Employees are entitled to 12 days of paid sick leave per year. A medical certificate is required for absences longer than 3 consecutive days.",
            "# Maternity Leave",
            "Employees are entitled to 180 days of paid maternity leave. Leave must commence no earlier than 45 days before the expected due date.",
            "# Paternity Leave",
            "Employees are entitled to 10 days of paid paternity leave, to be taken within 90 days of the child's birth.",
            "# Work From Home",
            "Employees may work from home up to 3 days per week, subject to manager approval. This section is new in version 2.",
            "# Travel Reimbursement",
            "Business travel expenses are reimbursed up to a cap of 800 USD per trip. Original receipts must be submitted within 30 days.",
        ],
    )

    write(
        "InsurancePolicy.docx",
        "Employee Insurance Policy",
        [
            "# Health Insurance",
            "The company provides group health insurance covering the employee, spouse, and up to two children. The annual coverage limit is 50,000 USD per family.",
            "# Dental and Vision",
            "Dental coverage is capped at 1,500 USD per year. Vision coverage includes one eye exam and one pair of prescription glasses per year.",
            "# Life Insurance",
            "Each employee is covered by a life insurance policy equal to 3 times their annual base salary.",
            "# Claims",
            "Claims must be submitted within 90 days of the date of service through the HR portal.",
        ],
    )

    write(
        "TravelPolicy.docx",
        "Corporate Travel Policy",
        [
            "# Booking",
            "All flights must be booked through the company's approved travel agency at least 14 days in advance where possible. Economy class is standard for flights under 6 hours.",
            "# Accommodation",
            "Hotel stays are reimbursed up to 200 USD per night in standard cities and 300 USD per night in designated high-cost cities.",
            "# Meals",
            "A daily meal allowance of 60 USD applies during business travel. Alcohol is not reimbursable.",
            "# Ground Transport",
            "Ride-share and taxi expenses are reimbursable. Rental cars require prior manager approval.",
        ],
    )

    print("\nDone. Upload these via http://127.0.0.1:8000/docs or curl.")
    return 0


if __name__ == "__main__":
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    raise SystemExit(main())
